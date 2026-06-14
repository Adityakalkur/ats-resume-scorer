"""
parser.py — Resume text extraction for PDF and DOCX files.

Returns a ParseResult dataclass containing:
  - full_text: str         raw extracted text (preserves line breaks)
  - sections: dict[str, str]  heading → content mapping (best-effort)
  - word_count: int
  - page_count: int
  - file_type: str         "pdf" | "docx"
  - warnings: list[str]   extraction issues detected
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import IO

from keywords import SECTION_HEADINGS, ALL_VALID_HEADINGS


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------

@dataclass
class ParseResult:
    full_text: str = ""
    sections: dict[str, str] = field(default_factory=dict)
    word_count: int = 0
    page_count: int = 0
    file_type: str = ""
    warnings: list[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def parse_resume(file_bytes: bytes, filename: str) -> ParseResult:
    """
    Entry point. Detect file type from filename extension, dispatch to the
    appropriate extractor, then run section segmentation on the raw text.
    Supported: PDF, DOCX/DOC, TXT, RTF, HTML/HTM
    """
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        result = _parse_pdf(file_bytes)
    elif ext in ("doc", "docx"):
        result = _parse_docx(file_bytes)
    elif ext == "txt":
        result = _parse_txt(file_bytes)
    elif ext == "rtf":
        result = _parse_rtf(file_bytes)
    elif ext in ("html", "htm"):
        result = _parse_html_resume(file_bytes)
    else:
        result = ParseResult(warnings=[
            f"Unsupported file type: .{ext}. Supported: PDF, DOCX, TXT, RTF, HTML"
        ])

    if result.full_text:
        result.sections = _segment_sections(result.full_text)
        result.word_count = len(result.full_text.split())

    return result


# ---------------------------------------------------------------------------
# PDF extraction
# ---------------------------------------------------------------------------

def _parse_pdf(file_bytes: bytes) -> ParseResult:
    """Extract text from PDF using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        return ParseResult(warnings=["pdfplumber is not installed. Run: pip install pdfplumber"])

    result = ParseResult(file_type="pdf")
    lines: list[str] = []

    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            result.page_count = len(pdf.pages)
            for page in pdf.pages:
                text = page.extract_text(x_tolerance=2, y_tolerance=2)
                if text:
                    lines.append(text)
                else:
                    result.warnings.append(
                        f"Page {page.page_number}: no text extracted — may contain images/scanned content."
                    )
    except Exception as exc:
        result.warnings.append(f"PDF parsing error: {exc}")
        return result

    result.full_text = "\n".join(lines)

    if not result.full_text.strip():
        result.warnings.append(
            "No readable text found. Your PDF may be image-based (scanned). "
            "ATS systems cannot read scanned resumes."
        )

    return result


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------

def _parse_docx(file_bytes: bytes) -> ParseResult:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
    except ImportError:
        return ParseResult(warnings=["python-docx is not installed. Run: pip install python-docx"])

    result = ParseResult(file_type="docx", page_count=1)
    paragraphs: list[str] = []

    try:
        doc = Document(io.BytesIO(file_bytes))

        # Check for tables (formatting red flag)
        if doc.tables:
            result.warnings.append(
                f"Detected {len(doc.tables)} table(s) in DOCX. "
                "Tables can confuse ATS parsers — consider switching to plain paragraphs."
            )

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract text from tables so we don't miss content
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in paragraphs:
                        paragraphs.append(cell_text)

    except Exception as exc:
        result.warnings.append(f"DOCX parsing error: {exc}")
        return result

    result.full_text = "\n".join(paragraphs)

    if not result.full_text.strip():
        result.warnings.append("No readable text found in DOCX.")

    return result


# ---------------------------------------------------------------------------
# Section segmentation (heuristic, works on both PDF and DOCX output)
# ---------------------------------------------------------------------------

_HEADING_RE = re.compile(
    r"^(?P<heading>[A-Z][A-Z &/\-]{2,40})$|"   # ALL CAPS heading
    r"^(?P<heading2>[A-Z][a-zA-Z &/\-]{2,40}):?\s*$",  # Title Case heading alone on line
    re.MULTILINE,
)


def _segment_sections(text: str) -> dict[str, str]:
    """
    Split the resume text into named sections by detecting heading lines.
    Returns a dict mapping canonical section name → content string.
    Unknown headings are kept as-is.
    """
    lines = text.splitlines()
    sections: dict[str, list[str]] = {}
    current_heading = "header"  # content before first recognized heading
    sections[current_heading] = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            sections[current_heading].append("")
            continue

        canonical = _match_heading(stripped)
        if canonical:
            current_heading = canonical
            if canonical not in sections:
                sections[canonical] = []
        else:
            sections[current_heading].append(stripped)

    # Join each section's lines back into text
    return {k: "\n".join(v).strip() for k, v in sections.items() if v}


def _match_heading(line: str) -> str | None:
    """
    Return canonical section name if the line looks like a resume section heading.
    Returns None if it's regular content.
    """
    lower = line.lower().strip(": ")

    # Direct match against known headings
    for canonical, variants in SECTION_HEADINGS.items():
        if lower in variants:
            return canonical

    # Fuzzy: if short line is mostly uppercase and vaguely matches a keyword
    if len(line) < 50 and line.isupper():
        for canonical, variants in SECTION_HEADINGS.items():
            for v in variants:
                if v in lower or lower in v:
                    return canonical

    return None


# ---------------------------------------------------------------------------
# Utility helpers used by scorer.py
# ---------------------------------------------------------------------------

def extract_bullets(section_text: str) -> list[str]:
    """
    Return individual bullet point strings from a section's text.
    Strips leading bullet markers (•, -, *, ▪, etc.) and blank lines.
    """
    bullet_re = re.compile(r"^[\•\-\*\▪\▸\➤\✓\–\—]\s*")
    bullets = []
    for line in section_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        cleaned = bullet_re.sub("", stripped).strip()
        if cleaned:
            bullets.append(cleaned)
    return bullets


def extract_all_bullets(sections: dict) -> list[str]:
    """
    Extract bullets from all sections in a sections dict.
    Convenience wrapper over extract_bullets for scorer.py.
    """
    all_bullets: list[str] = []
    for section_text in sections.values():
        all_bullets.extend(extract_bullets(section_text))
    return all_bullets


# ---------------------------------------------------------------------------
# Plain-text extraction
# ---------------------------------------------------------------------------

def _parse_txt(file_bytes: bytes) -> ParseResult:
    """Read plain-text resume (.txt)."""
    result = ParseResult(file_type="txt", page_count=1)
    try:
        text = file_bytes.decode("utf-8", errors="replace")
        result.full_text = text
        if not text.strip():
            result.warnings.append("The TXT file appears to be empty.")
    except Exception as exc:
        result.warnings.append(f"TXT parsing error: {exc}")
    return result


# ---------------------------------------------------------------------------
# RTF extraction
# ---------------------------------------------------------------------------

def _parse_rtf(file_bytes: bytes) -> ParseResult:
    """Extract text from RTF using striprtf."""
    result = ParseResult(file_type="rtf", page_count=1)
    try:
        from striprtf.striprtf import rtf_to_text
        raw = file_bytes.decode("latin-1", errors="replace")
        text = rtf_to_text(raw)
        result.full_text = "\n".join(
            line for line in text.splitlines() if line.strip()
        )
        if not result.full_text.strip():
            result.warnings.append("No readable text found in RTF.")
    except ImportError:
        result.warnings.append(
            "striprtf is not installed. Run: pip install striprtf"
        )
    except Exception as exc:
        result.warnings.append(f"RTF parsing error: {exc}")
    return result


# ---------------------------------------------------------------------------
# HTML extraction
# ---------------------------------------------------------------------------

def _parse_html_resume(file_bytes: bytes) -> ParseResult:
    """Extract text from HTML resume using BeautifulSoup."""
    result = ParseResult(file_type="html", page_count=1)
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(
            file_bytes.decode("utf-8", errors="replace"), "html.parser"
        )
        # Remove script / style noise
        for tag in soup(["script", "style", "meta", "head"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
        result.full_text = "\n".join(
            line.strip() for line in text.splitlines() if line.strip()
        )
        if not result.full_text.strip():
            result.warnings.append("No readable text found in HTML.")
    except ImportError:
        result.warnings.append(
            "beautifulsoup4 is not installed. Run: pip install beautifulsoup4"
        )
    except Exception as exc:
        result.warnings.append(f"HTML parsing error: {exc}")
    return result
